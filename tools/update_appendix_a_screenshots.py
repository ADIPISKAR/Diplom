from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


WORD_DIR = Path("word")
SOURCE = next(
    path
    for path in WORD_DIR.glob("*.docx")
    if "4.1" in path.name and "скриншот" in path.name and not path.name.startswith("~$")
)
OUTPUT = WORD_DIR / "Квалификационная работа - приложение А обновлено.docx"
SCREENSHOTS = WORD_DIR / "screenshots_4_1"

APPENDIX_TITLE_OLD = "Приложение A Листинги кода"
APPENDIX_TITLE_NEW = "Приложение A Скриншоты интерфейса"

IMAGES = [
    ("admin-dashboard.png", "Рисунок А.1 – Главная панель администратора"),
    ("admin-users.png", "Рисунок А.2 – Страница управления пользователями"),
    ("admin-categories.png", "Рисунок А.3 – Страница управления категориями оборудования"),
    ("admin-locations.png", "Рисунок А.4 – Страница управления местами хранения"),
    ("admin-equipment.png", "Рисунок А.5 – Страница реестра оборудования"),
    ("admin-equipment-form.png", "Рисунок А.6 – Форма добавления и редактирования оборудования"),
    ("admin-requests.png", "Рисунок А.7 – Страница заявок на оборудование"),
    ("employee-dashboard.png", "Рисунок А.8 – Панель сотрудника выдачи"),
    ("employee-requests.png", "Рисунок А.9 – Список заявок сотрудника выдачи"),
    ("employee-request-edit.png", "Рисунок А.10 – Страница обработки заявки сотрудником выдачи"),
]


def set_run_font(paragraph, *, bold=False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold


def set_title(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    paragraph.add_run(text)
    set_run_font(paragraph, bold=True)


def set_caption(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    paragraph.add_run(text)
    set_run_font(paragraph)


def set_image_paragraph(paragraph, image_path: Path) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    paragraph.add_run().add_picture(str(image_path), width=Cm(15.8))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


doc = Document(SOURCE)
paragraphs = doc.paragraphs

# Update the manual table-of-contents line for Appendix A.
for paragraph in paragraphs:
    normalized = paragraph.text.replace("\xa0", " ").strip()
    if normalized.startswith(APPENDIX_TITLE_OLD):
        paragraph.clear()
        paragraph.add_run(f"{APPENDIX_TITLE_NEW}\t78")
        set_run_font(paragraph)
        break

appendix_start = None
for index, paragraph in enumerate(paragraphs):
    normalized = paragraph.text.replace("\xa0", " ").strip()
    if normalized == APPENDIX_TITLE_OLD:
        appendix_start = index
        break

if appendix_start is None:
    raise RuntimeError("Appendix A title not found")

for index in range(len(paragraphs) - 1, appendix_start, -1):
    remove_paragraph(paragraphs[index])

set_title(paragraphs[appendix_start], APPENDIX_TITLE_NEW)

for image_name, caption in IMAGES:
    image_path = SCREENSHOTS / image_name
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    image_paragraph = doc.add_paragraph()
    set_image_paragraph(image_paragraph, image_path)
    caption_paragraph = doc.add_paragraph()
    set_caption(caption_paragraph, caption)

doc.save(OUTPUT)
print(OUTPUT)
for _, caption in IMAGES:
    print(caption)
