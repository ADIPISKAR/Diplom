from pathlib import Path
from docx import Document


def show_doc(path: Path):
    doc = Document(path)
    print(f"DOC={path.name} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for i, p in enumerate(doc.paragraphs, 1):
        t = p.text.strip()
        if not t:
            continue
        if "Прилож" in t or "Листинг" in t or "public function" in t or "<?php" in t or i > len(doc.paragraphs) - 180:
            fmt = p.paragraph_format
            run = next((r for r in p.runs if r.text.strip()), None)
            font = run.font.name if run else None
            size = run.font.size.pt if run and run.font.size else None
            print(f"{i}: style={p.style.name} align={p.alignment} first={fmt.first_line_indent} left={fmt.left_indent} line={fmt.line_spacing} font={font} size={size}: {t[:180]}")


for file in sorted(Path("word").glob("*.docx")):
    if "Пример" in file.name:
        show_doc(file)
