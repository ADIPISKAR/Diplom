from pathlib import Path

from docx import Document
from docx.shared import Pt


WORD_DIR = Path("word")
OUTPUT = WORD_DIR / "Итоговый вариант - проверенные схемы.docx"


def select_source() -> Path:
    return max(
        [path for path in WORD_DIR.glob("*.docx") if not path.name.startswith("~$")],
        key=lambda path: path.stat().st_mtime,
    )


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def main() -> None:
    source = select_source()
    doc = Document(str(source.resolve()))

    # The document should contain five project schemes from "Таблицы.docx".
    # Remove the leftover old DFD zero-level image/caption and its explanation.
    for index in range(276, 269, -1):
        remove_paragraph(doc.paragraphs[index])

    set_text(
        doc.paragraphs[269],
        "На рисунке 3.2 представлена DFD-диаграмма потоков данных, раскрывающая "
        "взаимодействие пользователя, сотрудника выдачи, администратора, базы данных "
        "и основных процессов веб-приложения.",
    )

    # After removal, paragraph indexes shift by seven.
    replacements = {
        271: "Рисунок 3.2 – DFD-диаграмма потоков данных",
        288: "На рисунке 3.3 представлена контекстная диаграмма функциональной модели IDEF0.",
        290: "Рисунок 3.3 – Контекстная диаграмма функциональной модели",
        295: "На рисунке 3.4 представлена декомпозиция контекстной диаграммы функциональной модели.",
        297: "Рисунок 3.4 – Декомпозиция контекстной диаграммы",
        375: "Рисунок 3.5 – Структура данных и взаимосвязи между таблицами.",
    }

    for index, text in replacements.items():
        set_text(doc.paragraphs[index], text)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(source)


if __name__ == "__main__":
    main()
