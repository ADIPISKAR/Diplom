from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image


WORD_DIR = Path("word")
EXTRACTED = WORD_DIR / "extracted_tables"
OUTPUT = WORD_DIR / "Итоговый вариант - таблицы и источники.docx"


def select_final_doc() -> Path:
    candidates = [
        path
        for path in WORD_DIR.glob("*.docx")
        if not path.name.startswith("~$")
    ]
    # The final diploma is the large document with the appendix screenshots.
    return max(candidates, key=lambda path: path.stat().st_size)


def select_tables_doc() -> Path:
    for path in WORD_DIR.glob("*.docx"):
        if path.name.startswith("~$"):
            continue
        try:
            doc = Document(str(path.resolve()))
        except Exception:
            continue
        if len(doc.inline_shapes) == 5 and len(doc.paragraphs) <= 10:
            return path
    raise RuntimeError("Tables document not found")


def extract_table_images() -> list[Path]:
    tables_doc = select_tables_doc()
    EXTRACTED.mkdir(exist_ok=True)
    with ZipFile(tables_doc) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        media.sort(key=lambda name: int("".join(filter(str.isdigit, Path(name).stem)) or 0))
        paths = []
        for index, name in enumerate(media, 1):
            target = EXTRACTED / f"table_image_{index}{Path(name).suffix}"
            target.write_bytes(archive.read(name))
            paths.append(target)
    return paths


def clear_paragraph(paragraph) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5


def add_image_to_paragraph(paragraph, image_path: Path, width_cm: float = 15.8) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))


def insert_paragraph_before(paragraph):
    new_p = paragraph._element.addprevious(paragraph._element.__class__())
    # addprevious returns None for lxml elements; fetch inserted sibling.
    inserted = paragraph._element.getprevious()
    return paragraph._parent._paragraph_cls(inserted, paragraph._parent)


def find_caption(doc: Document, prefix: str):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index, paragraph
    raise RuntimeError(f"Caption not found: {prefix}")


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def source_sort_key(text: str):
    normalized = text.strip().lower().replace("ё", "е")
    first = normalized[:1]
    if "а" <= first <= "я":
        group = 0
    elif "a" <= first <= "z":
        group = 1
    else:
        group = 2
    return (group, normalized)


def replace_sources_with_sorted(doc: Document) -> None:
    heading = "Перечень использованных информационных ресурсов"
    appendix = "Приложение A"
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == heading)
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if p.text.strip().startswith(appendix))

    source_paragraphs = doc.paragraphs[start + 1 : end]
    sources = [p.text for p in source_paragraphs if p.text.strip()]
    sorted_sources = sorted(sources, key=source_sort_key)

    template = source_paragraphs[0]
    for paragraph, text in zip(source_paragraphs, sorted_sources):
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        paragraph.alignment = template.alignment
        paragraph.paragraph_format.left_indent = template.paragraph_format.left_indent
        paragraph.paragraph_format.right_indent = template.paragraph_format.right_indent
        paragraph.paragraph_format.first_line_indent = template.paragraph_format.first_line_indent
        paragraph.paragraph_format.space_before = template.paragraph_format.space_before
        paragraph.paragraph_format.space_after = template.paragraph_format.space_after
        paragraph.paragraph_format.line_spacing = template.paragraph_format.line_spacing


def main() -> None:
    final_doc = select_final_doc()
    images = extract_table_images()
    process_scheme = images[-2]
    db_scheme = images[-1]

    # Make sure the files are readable image files before inserting.
    Image.open(process_scheme).verify()
    Image.open(db_scheme).verify()

    doc = Document(str(final_doc.resolve()))

    # Figure 3.5 has a caption but no actual image before it.
    fig35_index, fig35_caption = find_caption(doc, "Рисунок 3.5")
    previous = doc.paragraphs[fig35_index - 1] if fig35_index > 0 else None
    if previous is not None and not previous.text.strip() and not paragraph_has_drawing(previous):
        target = previous
    else:
        target = fig35_caption.insert_paragraph_before()
    add_image_to_paragraph(target, process_scheme)

    # Figure 3.6 already has an image paragraph before the caption; replace it.
    fig36_index, fig36_caption = find_caption(doc, "Рисунок 3.6")
    if fig36_index > 0 and paragraph_has_drawing(doc.paragraphs[fig36_index - 1]):
        add_image_to_paragraph(doc.paragraphs[fig36_index - 1], db_scheme)
    else:
        target = fig36_caption.insert_paragraph_before()
        add_image_to_paragraph(target, db_scheme)

    replace_sources_with_sorted(doc)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(process_scheme)
    print(db_scheme)


if __name__ == "__main__":
    main()
