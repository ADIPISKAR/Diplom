from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


WORD_DIR = Path("word")
SOURCE = next(
    path
    for path in WORD_DIR.glob("*.docx")
    if "4.1" in path.name and "обновлена" in path.name and not path.name.startswith("~$")
)
OUTPUT = WORD_DIR / "Квалификационная работа - глава 4.1 со скриншотами.docx"
SCREENSHOTS = WORD_DIR / "screenshots_4_1"

CAPTION_TO_IMAGE = {
    "Рисунок 4.3": "admin-dashboard.png",
    "Рисунок 4.4": "admin-users.png",
    "Рисунок 4.5": "admin-categories.png",
    "Рисунок 4.6": "admin-locations.png",
    "Рисунок 4.7": "admin-equipment.png",
    "Рисунок 4.8": "admin-equipment-form.png",
    "Рисунок 4.9": "admin-requests.png",
    "Рисунок 4.10": "employee-dashboard.png",
    "Рисунок 4.11": "employee-request-edit.png",
}


def has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def clear_paragraph(paragraph) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


doc = Document(SOURCE)
paragraphs = doc.paragraphs
replaced = []

for index, paragraph in enumerate(paragraphs):
    caption = paragraph.text.strip()
    key = next((prefix for prefix in CAPTION_TO_IMAGE if caption.startswith(prefix)), None)
    if key is None:
        continue

    image_path = SCREENSHOTS / CAPTION_TO_IMAGE[key]
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    target = paragraphs[index - 1] if index > 0 and has_drawing(paragraphs[index - 1]) else None
    if target is None and index > 1 and not paragraphs[index - 1].text.strip():
        target = paragraphs[index - 1]
    if target is None:
        target = paragraph.insert_paragraph_before()

    clear_paragraph(target)
    run = target.add_run()
    run.add_picture(str(image_path), width=Cm(15.8))
    replaced.append(f"{key}: {image_path.name}")

doc.save(OUTPUT)
print(OUTPUT)
print("\n".join(replaced))
